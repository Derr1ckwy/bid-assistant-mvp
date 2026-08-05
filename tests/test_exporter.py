from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from PIL import Image

from bid_assistant.exporter import export_docx
from bid_assistant.models import (
    ChapterDraft,
    ProjectInfo,
    RequirementItem,
    ReviewIssue,
    ReviewReport,
    ScoringItem,
    TenderAnalysis,
)


def test_export_docx_can_be_reopened(tmp_path: Path) -> None:
    output = tmp_path / "draft.docx"
    analysis = TenderAnalysis(
        project_info=ProjectInfo(
            project_name="智慧园区项目",
            purchaser="测试采购单位",
            budget="180 万元",
            bid_deadline="2026-09-15 09:30",
        ),
        mandatory_requirements=[
            RequirementItem(content="必须提交营业执照。", source_page=3, status="已确认")
        ],
        scoring_items=[
            ScoringItem(criterion="技术方案完整性", points="20", source_page=8)
        ],
    )
    drafts = [
        ChapterDraft(
            chapter_id="chapter_1",
            title="技术方案",
            markdown="## 总体设计\n\n1. 建立统一平台。\n\n- 支持人工复核。",
        )
    ]
    review = ReviewReport(
        issues=[
            ReviewIssue(
                severity="高",
                category="废标风险",
                message="必须复核营业执照。",
                suggestion="确认文件有效期。",
            )
        ]
    )

    result = export_docx(output, analysis, drafts, review)
    reopened = Document(result)
    text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    footer_xml = reopened.sections[0].footer._element.xml
    first_header_xml = reopened.sections[0].first_page_header._element.xml
    first_footer_xml = reopened.sections[0].first_page_footer._element.xml

    assert result.exists()
    assert "智慧园区项目" in text
    assert "第 1 章 技术方案" in text
    assert "内部核对事项" not in text
    assert "内部复核报告" not in text
    assert "PAGE" in footer_xml
    assert reopened.sections[0].different_first_page_header_footer is True
    assert "智慧园区项目" not in first_header_xml
    assert "PAGE" not in first_footer_xml
    assert len(reopened.sections[0].footer.tables) == 1
    footer_table = reopened.sections[0].footer.tables[0]
    assert "AI" not in footer_table.rows[0].cells[0].text
    assert footer_table.rows[0].cells[1].paragraphs[0].alignment == 1
    assert len(reopened.sections[0].header.tables) == 1
    header_table = reopened.sections[0].header.tables[0]
    assert header_table.rows[0].cells[0].text == "智慧园区项目"
    assert header_table.rows[0].cells[1].paragraphs[0].alignment == 2

    section = reopened.sections[0]
    assert abs(section.page_width - Cm(21)) < 1000
    assert abs(section.page_height - Cm(29.7)) < 1000

    normal = reopened.styles["Normal"]
    assert normal.font.size.pt == 12
    assert normal.paragraph_format.line_spacing.pt == 26
    assert normal._element.xml.find('w:eastAsia="宋体"') >= 0

    cover_title = next(paragraph for paragraph in reopened.paragraphs if paragraph.text == "智慧园区项目")
    assert cover_title.runs[0]._element.find(qn("w:rPr")).find(qn("w:spacing")) is not None

    cover_metadata = reopened.tables[0]
    assert [cell.width for cell in cover_metadata.rows[0].cells] == [2750 * 635, 6151 * 635]
    assert cover_metadata.rows[0].cells[0]._tc.tcPr.find(qn("w:noWrap")) is not None

    assert len(reopened.tables) == 1
    for table in reopened.tables:
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        width = table._tbl.tblPr.find(qn("w:tblW"))
        assert layout is not None and layout.get(qn("w:type")) == "fixed"
        assert width is not None and width.get(qn("w:type")) == "dxa"


def test_export_uses_real_template_and_appends_qualification_images(tmp_path: Path) -> None:
    template = tmp_path / "客户模板.docx"
    template_document = Document()
    header = template_document.sections[0].header.paragraphs[0]
    header.add_run("{{PROJECT_").bold = True
    header.add_run("NAME}}")
    header.add_run(" | ")
    header.add_run("{{PURCHASER}}")
    template_document.add_paragraph("项目预算：{{BUDGET}}")
    template_document.add_paragraph("{{BID_CONTENT}}")
    template_document.save(template)
    image_path = tmp_path / "信息安全认证.png"
    Image.new("RGB", (640, 420), "white").save(image_path)
    output = tmp_path / "template-output.docx"
    analysis = TenderAnalysis(project_info=ProjectInfo(project_name="模板适配项目", purchaser="采购人", budget="100 万元"))
    drafts = [ChapterDraft(chapter_id="c1", title="实施方案", markdown="## 项目理解\n\n正文内容。")]

    export_docx(output, analysis, drafts, ReviewReport(), template_path=template, qualification_images=[image_path])
    reopened = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)

    assert "项目预算：100 万元" in body_text
    assert "{{BID_CONTENT}}" not in body_text
    assert "第 1 章 实施方案" in body_text
    assert "资质证明材料" in body_text
    assert "模板适配项目 | 采购人" in reopened.sections[0].header.paragraphs[0].text
    assert reopened.sections[0].header.paragraphs[0].runs[0].bold is True
    assert len(reopened.inline_shapes) == 1


def test_export_cleans_model_markdown_and_uses_formal_monochrome_tables(tmp_path: Path) -> None:
    output = tmp_path / "formal.docx"
    analysis = TenderAnalysis(project_info=ProjectInfo(project_name="正式排版测试项目"))
    drafts = [
        ChapterDraft(
            chapter_id="c1",
            title="技术方案",
            markdown=(
                "```markdown\n---\ndoc_id: TEST-001\nsimulation: true\n---\n"
                "# 技术方案\n\n## 实施步骤\n\n"
                "1. 第一项。[资料1]\n2. 第二项。\n\n正文分隔。\n\n"
                "1. 新的一组。\n2. 新的第二项。\n\n"
                "| 序号 | 工作内容 | 状态 |\n| --- | --- | --- |\n"
                "| 1 | 完成现场复核与方案确认 | 已完成 |\n```"
            ),
        )
    ]

    export_docx(output, analysis, drafts)
    reopened = Document(output)
    text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    footer_text = "\n".join(
        cell.text
        for table in reopened.sections[0].footer.tables
        for row in table.rows
        for cell in row.cells
    )

    assert "doc_id" not in text
    assert "simulation:" not in text
    assert "```" not in text
    assert "[资料1]" not in text
    assert sum(paragraph.text == "技术方案" for paragraph in reopened.paragraphs) == 0
    assert "AI" not in footer_text

    numbered = [
        paragraph
        for paragraph in reopened.paragraphs
        if paragraph.text in {"第一项。", "第二项。", "新的一组。", "新的第二项。"}
    ]
    num_ids = [
        paragraph._p.pPr.numPr.numId.val
        for paragraph in numbered
    ]
    assert num_ids[0] == num_ids[1]
    assert num_ids[2] == num_ids[3]
    assert num_ids[0] != num_ids[2]

    content_table = reopened.tables[-1]
    header_fill = content_table.rows[0].cells[0]._tc.tcPr.find(qn("w:shd")).get(qn("w:fill"))
    header_color = content_table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb
    assert header_fill == "F2F2F2"
    assert str(header_color) == "000000"


def test_internal_appendices_are_opt_in(tmp_path: Path) -> None:
    output = tmp_path / "internal.docx"
    analysis = TenderAnalysis(
        project_info=ProjectInfo(project_name="内部评审项目"),
        mandatory_requirements=[RequirementItem(content="必须提交营业执照。")],
    )
    drafts = [ChapterDraft(chapter_id="c1", title="响应方案", markdown="## 项目响应\n\n正文。")]
    review = ReviewReport(
        issues=[ReviewIssue(severity="高", category="资格", message="待确认。", suggestion="复核。")]
    )

    export_docx(output, analysis, drafts, review, include_internal_appendices=True)
    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)

    assert "内部核对事项" in text
    assert "内部复核报告" in text
