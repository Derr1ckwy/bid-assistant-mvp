from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from bid_assistant.parsers import (
    KNOWLEDGE_EXTENSIONS,
    TENDER_EXTENSIONS,
    DocumentParseError,
    parse_document,
    parse_document_bytes,
)
from bid_assistant.models import ParsedDocument, ParsedPage


def test_parse_utf8_text(tmp_path: Path) -> None:
    source = tmp_path / "tender.txt"
    source.write_text("项目名称：测试项目\n投标人必须提交材料。", encoding="utf-8")

    parsed = parse_document(source)

    assert parsed.file_type == "txt"
    assert parsed.char_count > 10
    assert "测试项目" in parsed.full_text
    assert parsed.pages[0].page_number == 1


def test_parse_gb18030_text(tmp_path: Path) -> None:
    source = tmp_path / "legacy.txt"
    source.write_bytes("采购人：测试单位".encode("gb18030"))

    parsed = parse_document(source)

    assert "测试单位" in parsed.full_text


def test_parse_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    source = tmp_path / "tender.docx"
    document = Document()
    document.add_paragraph("项目名称：DOCX 测试项目")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "评分项"
    table.cell(0, 1).text = "10 分"
    document.save(source)

    parsed = parse_document(source)

    assert "DOCX 测试项目" in parsed.full_text
    assert "评分项 | 10 分" in parsed.full_text


def test_blank_pdf_is_marked_as_possible_scan(tmp_path: Path) -> None:
    source = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with source.open("wb") as stream:
        writer.write(stream)

    parsed = parse_document(source)

    assert parsed.possible_scanned_document is True
    assert any("扫描件" in warning for warning in parsed.warnings)


def test_parse_csv_preserves_headers_and_row_values(tmp_path: Path) -> None:
    source = tmp_path / "company.csv"
    source.write_text(
        "公司名称,资质名称,有效期\n武汉灵坤数字科技有限公司,ISO 9001,2028-12-31\n",
        encoding="utf-8-sig",
    )

    parsed = parse_document(source)

    assert parsed.file_type == "csv"
    assert "公司名称: 武汉灵坤数字科技有限公司" in parsed.full_text
    assert "资质名称: ISO 9001" in parsed.full_text
    assert "有效期: 2028-12-31" in parsed.full_text


def test_parse_gb18030_csv(tmp_path: Path) -> None:
    source = tmp_path / "legacy.csv"
    source.write_bytes("产品名称,功能\n投标助手,知识检索".encode("gb18030"))

    parsed = parse_document(source)

    assert "产品名称: 投标助手" in parsed.full_text
    assert "功能: 知识检索" in parsed.full_text


def test_parse_nested_json_preserves_field_paths(tmp_path: Path) -> None:
    source = tmp_path / "company.json"
    source.write_text(
        '{"company":{"name":"武汉灵坤","certificates":[{"name":"ISO 9001","valid":true}]}}',
        encoding="utf-8",
    )

    parsed = parse_document(source)

    assert parsed.file_type == "json"
    assert "company.name: 武汉灵坤" in parsed.full_text
    assert "company.certificates[0].name: ISO 9001" in parsed.full_text
    assert "company.certificates[0].valid: true" in parsed.full_text


def test_invalid_json_raises_readable_error(tmp_path: Path) -> None:
    source = tmp_path / "broken.json"
    source.write_text('{"company":', encoding="utf-8")

    with pytest.raises(DocumentParseError, match="JSON 无法读取"):
        parse_document(source)


def test_csv_and_json_are_limited_to_knowledge_uploads() -> None:
    assert ".csv" in KNOWLEDGE_EXTENSIONS
    assert ".json" in KNOWLEDGE_EXTENSIONS
    assert ".csv" not in TENDER_EXTENSIONS
    assert ".json" not in TENDER_EXTENSIONS


def test_parse_document_bytes_validates_uploaded_json() -> None:
    parsed = parse_document_bytes(
        "company.json",
        '{"company":{"name":"武汉灵坤"}}'.encode("utf-8"),
    )

    assert "company.name: 武汉灵坤" in parsed.full_text


def test_auto_parser_uses_mineru_for_scanned_pdf(tmp_path: Path) -> None:
    source = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with source.open("wb") as target:
        writer.write(target)

    class FakeMinerU:
        def parse(self, path):
            text = "OCR 提取的招标内容" * 20
            return ParsedDocument(
                filename=Path(path).name,
                file_type="pdf",
                pages=[ParsedPage(page_number=1, text=text)],
                full_text=text,
                char_count=len(text),
                parser_engine="mineru",
            )

    parsed = parse_document(source, mode="auto", mineru_client=FakeMinerU())

    assert parsed.parser_engine == "mineru"
    assert "OCR 提取" in parsed.full_text


def test_mineru_failure_falls_back_to_native_parser(tmp_path: Path) -> None:
    source = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with source.open("wb") as target:
        writer.write(target)

    class BrokenMinerU:
        def parse(self, path):
            raise RuntimeError("not installed")

    parsed = parse_document(source, mode="mineru", mineru_client=BrokenMinerU())

    assert parsed.parser_engine == "native"
    assert any("已保留原生结果" in warning for warning in parsed.warnings)
