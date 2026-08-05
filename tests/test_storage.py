import io
import json
import sqlite3
import zipfile
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from bid_assistant.storage import (
    MAX_DRAFT_VERSIONS,
    MAX_EXPORT_VERSIONS,
    MAX_PACKAGE_VERSIONS,
    ProjectArchiveError,
    ProjectStore,
    format_beijing_time,
    safe_filename,
)


def test_project_store_round_trip(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("测试项目")

    source = store.save_source(project["id"], "招标?.txt", "内容".encode("utf-8"))
    store.save_json(project["id"], "analysis", {"ok": True})
    knowledge = store.save_knowledge_file(project["id"], "company", "企业资料.txt", b"company")

    assert source.exists()
    assert source.name == "招标_.txt"
    assert store.source_path(project["id"]) == source
    assert store.load_json(project["id"], "analysis") == {"ok": True}
    assert knowledge in store.list_knowledge_files(project["id"])["company"]
    assert store.get_project(project["id"])["status"] == "knowledge_ready"


def test_store_rejects_invalid_identifiers(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")

    with pytest.raises(ValueError):
        store.project_dir("../escape")
    with pytest.raises(ValueError):
        store.save_knowledge_file("project", "unknown", "a.txt", b"x")
    with pytest.raises(ValueError):
        store.save_attachment_file("project", "unknown", "a.txt", b"x")


@pytest.mark.parametrize("category", ["company", "product", "history"])
def test_knowledge_files_are_categorized_and_deletable(tmp_path: Path, category: str) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("知识资料删除测试")
    source = store.save_knowledge_file(project["id"], category, "产品参数?.json", b"{}")
    reference = f"{category}/{source.name}"

    assert source.name == "产品参数_.json"
    assert store.knowledge_path(project["id"], reference) == source
    assert store.knowledge_path(project["id"], "../bad") is None
    assert store.knowledge_path(project["id"], "unknown/file.json") is None
    assert store.delete_knowledge_file(project["id"], reference) is True
    assert store.delete_knowledge_file(project["id"], reference) is False
    assert store.list_knowledge_files(project["id"])[category] == []


def test_word_template_and_qualification_images_are_managed_per_project(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("模板测试")
    template_buffer = io.BytesIO()
    template_document = Document()
    template_document.add_paragraph("{{PROJECT_NAME}}")
    template_document.add_paragraph("{{BID_CONTENT}}")
    template_document.save(template_buffer)
    template = store.save_word_template(project["id"], "甲方模板.docx", template_buffer.getvalue())

    image_buffer = io.BytesIO()
    Image.new("RGB", (320, 180), "white").save(image_buffer, format="PNG")
    image = store.save_qualification_image(project["id"], "营业执照.png", image_buffer.getvalue())

    assert store.word_template_path(project["id"]) == template
    assert store.list_qualification_images(project["id"]) == [image]
    assert store.delete_qualification_image(project["id"], image.name) is True
    assert store.delete_word_template(project["id"]) is True


def test_invalid_template_and_image_are_rejected(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("非法文件测试")

    with pytest.raises(ValueError):
        store.save_word_template(project["id"], "模板.docx", b"not-docx")
    with pytest.raises(ValueError):
        store.save_qualification_image(project["id"], "证书.png", b"not-image")


def test_safe_filename_removes_path_and_windows_characters() -> None:
    assert safe_filename("..\\folder\\bad?.txt") == "bad_.txt"
    assert safe_filename("../folder/bad?.txt") == "bad_.txt"


def test_existing_database_is_migrated_for_archiving(tmp_path: Path) -> None:
    data_dir = tmp_path / "data"
    data_dir.mkdir()
    with sqlite3.connect(data_dir / "app.db") as conn:
        conn.execute(
            """
            CREATE TABLE projects (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                status TEXT NOT NULL DEFAULT 'new',
                source_filename TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
            """
        )
        conn.execute(
            "INSERT INTO projects VALUES ('proj_old', '旧项目', 'new', NULL, '2026-01-01', '2026-01-01')"
        )

    store = ProjectStore(data_dir)

    assert store.get_project("proj_old")["archived"] == 0


def test_archive_and_restore_project(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("归档测试")

    store.set_project_archived(project["id"], True)
    assert store.list_projects() == []
    assert store.list_projects(include_archived=True)[0]["archived"] == 1

    store.set_project_archived(project["id"], False)
    assert store.list_projects()[0]["id"] == project["id"]


def test_duplicate_project_copies_business_data_without_outputs_or_versions(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("原项目")
    project_id = project["id"]
    store.save_source(project_id, "招标文件.txt", b"tender")
    store.save_json(project_id, "analysis", {"ok": True})
    store.save_json(project_id, "analysis_baseline", {"baseline": True})
    store.save_json(
        project_id,
        "analysis_baseline_meta",
        {"origin": "自动分析", "source_fingerprint": "abc"},
    )
    store.save_json(project_id, "analysis_acceptance", {"complete": True})
    store.save_json(project_id, "drafts", [{"title": "第一版"}])
    store.save_json(project_id, "review", {"issues": []})
    store.save_knowledge_file(project_id, "company", "企业资料.txt", b"company")
    template_buffer = io.BytesIO()
    Document().save(template_buffer)
    store.save_word_template(project_id, "公司模板.docx", template_buffer.getvalue())
    image_buffer = io.BytesIO()
    Image.new("RGB", (100, 80), "white").save(image_buffer, format="PNG")
    store.save_qualification_image(project_id, "营业执照.png", image_buffer.getvalue())
    store.save_attachment_file(project_id, "qualification", "营业执照.pdf", b"license")
    store.save_json(project_id, "submission_checklist", [{"name": "营业执照", "status": "已备妥"}])
    store.save_draft_version(project_id, [{"title": "第一版"}], "首次保存")
    old_output = store.output_path(project_id, "旧初稿.docx")
    old_output.write_bytes(b"docx")
    store.record_export_version(
        project_id,
        old_output,
        version=1,
        chapter_count=1,
        review_summary={"pending": 0, "high": 0, "medium": 0, "low": 0},
    )
    old_package = store.package_path(project_id, "旧提交包_P001.zip")
    old_package.write_bytes(b"package")
    store.record_package_version(
        project_id,
        old_package,
        version=1,
        word_version={"version": 1, "filename": old_output.name},
        checklist_summary={"total": 1, "required": 1, "ready": 1, "complete": True},
        attachment_count=1,
    )
    store.update_project(project_id, status="packaged")

    duplicate = store.duplicate_project(project_id)

    assert duplicate["id"] != project_id
    assert duplicate["name"] == "原项目（副本）"
    assert duplicate["status"] == "review_generated"
    assert duplicate["archived"] == 0
    assert store.load_json(duplicate["id"], "analysis") == {"ok": True}
    assert store.load_json(duplicate["id"], "analysis_baseline") == {"baseline": True}
    assert store.load_json(duplicate["id"], "analysis_baseline_meta") == {
        "origin": "自动分析",
        "source_fingerprint": "abc",
    }
    assert store.load_json(duplicate["id"], "analysis_acceptance") == {"complete": True}
    assert store.load_json(duplicate["id"], "drafts") == [{"title": "第一版"}]
    assert store.list_knowledge_files(duplicate["id"])["company"][0].read_bytes() == b"company"
    assert store.word_template_path(duplicate["id"]).name == "公司模板.docx"
    assert store.list_qualification_images(duplicate["id"])[0].name == "营业执照.png"
    assert sum(len(paths) for paths in store.list_attachment_files(duplicate["id"]).values()) == 0
    assert store.load_json(duplicate["id"], "submission_checklist") is None
    assert not (store.project_dir(duplicate["id"]) / "output").exists()
    assert not (store.project_dir(duplicate["id"]) / "packages").exists()
    assert store.list_draft_versions(duplicate["id"]) == []
    assert store.list_export_versions(duplicate["id"]) == []
    assert store.list_package_versions(duplicate["id"]) == []


def test_draft_versions_can_be_listed_and_restored(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("版本测试")
    project_id = project["id"]
    first_drafts = [{"chapter_id": "a", "title": "第一章", "markdown": "第一版"}]
    second_drafts = [{"chapter_id": "a", "title": "第一章", "markdown": "第二版"}]

    store.save_json(project_id, "drafts", first_drafts)
    first_version = store.save_draft_version(project_id, first_drafts, "首次生成")
    store.save_json(project_id, "drafts", second_drafts)
    store.save_draft_version(project_id, second_drafts, "人工修改")
    store.save_json(project_id, "review", {"issues": [{"message": "旧复核"}]})

    versions = store.list_draft_versions(project_id)
    restored = store.restore_draft_version(project_id, first_version["id"])

    assert len(versions) == 2
    assert versions[0]["reason"] == "人工修改"
    assert restored == first_drafts
    assert store.load_json(project_id, "drafts") == first_drafts
    assert store.load_json(project_id, "review") is None
    assert len(store.list_draft_versions(project_id)) == 3
    assert any(item["reason"] == "恢复版本前自动快照" for item in store.list_draft_versions(project_id))

    with pytest.raises(ValueError):
        store.load_draft_version(project_id, "../bad")


def test_draft_version_retention_keeps_latest_versions(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("版本上限测试")
    drafts = [{"chapter_id": "a", "title": "第一章", "markdown": "正文"}]

    for index in range(MAX_DRAFT_VERSIONS + 2):
        store.save_draft_version(project["id"], drafts, f"版本 {index}")

    versions = store.list_draft_versions(project["id"])

    assert len(versions) == MAX_DRAFT_VERSIONS
    assert versions[0]["reason"] == f"版本 {MAX_DRAFT_VERSIONS + 1}"
    assert all(item["reason"] != "版本 0" for item in versions)


def test_project_progress_uses_persisted_artifacts(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("进度测试")
    project_id = project["id"]

    store.save_source(project_id, "招标文件.txt", b"tender")
    store.save_json(project_id, "parsed", {"full_text": "tender"})
    store.save_json(project_id, "analysis", {"project_info": {}})
    store.save_json(project_id, "drafts", [{"title": "draft"}])
    store.save_knowledge_file(project_id, "company", "企业资料.txt", b"company")
    store.save_attachment_file(project_id, "qualification", "营业执照.pdf", b"license")

    progress = store.project_progress(project_id)

    assert progress["completed"] == 4
    assert progress["percent"] == 50
    assert progress["knowledge_files"] == 1
    assert progress["attachment_files"] == 1

    store.save_json(
        project_id,
        "submission_checklist",
        [{"name": "营业执照", "required": True, "status": "已备妥"}],
    )

    completed_progress = store.project_progress(project_id)
    assert completed_progress["completed"] == 5
    assert completed_progress["steps"][-2]["complete"] is True

    package = store.package_path(project_id, "提交包_P001.zip")
    package.write_bytes(b"package")
    store.record_package_version(
        project_id,
        package,
        version=1,
        word_version={"version": 1, "filename": "投标文件_V001.docx"},
        checklist_summary={"total": 1, "required": 1, "ready": 1, "complete": True},
        attachment_count=1,
    )

    packaged_progress = store.project_progress(project_id)
    assert packaged_progress["completed"] == 6
    assert packaged_progress["steps"][-1] == {
        "key": "package",
        "label": "提交打包",
        "complete": True,
    }


def test_attachment_files_are_categorized_unique_and_deletable(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("附件测试")

    first = store.save_attachment_file(project["id"], "qualification", "营业执照?.pdf", b"first")
    second = store.save_attachment_file(project["id"], "qualification", "营业执照?.pdf", b"second")
    reference = f"qualification/{first.name}"

    assert first.name == "营业执照_.pdf"
    assert second.name == "营业执照_ (2).pdf"
    assert store.attachment_path(project["id"], reference) == first
    assert store.attachment_path(project["id"], "../bad") is None
    assert store.delete_attachment_file(project["id"], reference) is True
    assert store.delete_attachment_file(project["id"], reference) is False


def test_word_outputs_are_versioned_and_recorded(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("导出版本测试")

    first = store.next_output_version(project["id"], "测试项目_投标文件初稿.docx")
    first["path"].write_bytes(b"first docx")
    first_record = store.record_export_version(
        project["id"],
        first["path"],
        version=first["version"],
        chapter_count=6,
        review_summary={"pending": 3, "high": 1, "medium": 2, "low": 0},
        warning_count=2,
        note="第一次评审",
        template_filename="甲方模板.docx",
        qualification_image_count=3,
    )
    second = store.next_output_version(project["id"], "测试项目_投标文件初稿.docx")
    second["path"].write_bytes(b"second docx")
    store.record_export_version(
        project["id"],
        second["path"],
        version=second["version"],
        chapter_count=6,
        review_summary={"pending": 0, "high": 0, "medium": 0, "low": 0},
    )

    versions = store.list_export_versions(project["id"])

    assert first["filename"].endswith("_V001.docx")
    assert second["filename"].endswith("_V002.docx")
    assert first_record["sha256"]
    assert first_record["created_at"].endswith("+08:00")
    assert [item["version"] for item in versions] == [2, 1]
    assert versions[1]["note"] == "第一次评审"
    assert versions[1]["review_summary"]["high"] == 1
    assert versions[1]["template_filename"] == "甲方模板.docx"
    assert versions[1]["qualification_image_count"] == 3


def test_word_version_time_is_displayed_in_beijing_time() -> None:
    assert format_beijing_time("2026-08-05T03:40:06.349+00:00") == "2026-08-05 11:40:06"
    assert format_beijing_time("2026-08-05T11:40:06.349+08:00") == "2026-08-05 11:40:06"


def test_word_output_retention_keeps_latest_versions(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("导出版本上限测试")

    for index in range(MAX_EXPORT_VERSIONS + 2):
        target = store.next_output_version(project["id"], "投标文件初稿.docx")
        target["path"].write_bytes(f"docx-{index}".encode())
        store.record_export_version(
            project["id"],
            target["path"],
            version=target["version"],
            chapter_count=1,
            review_summary={"pending": 0, "high": 0, "medium": 0, "low": 0},
        )

    versions = store.list_export_versions(project["id"])
    output_files = list((store.project_dir(project["id"]) / "output").glob("*.docx"))

    assert len(versions) == MAX_EXPORT_VERSIONS
    assert len(output_files) == MAX_EXPORT_VERSIONS
    assert versions[0]["version"] == MAX_EXPORT_VERSIONS + 2
    assert all(item["version"] != 1 for item in versions)


def test_submission_packages_are_versioned_and_recorded(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("提交包版本测试")

    first = store.next_package_version(project["id"], "测试项目_最终提交包.zip")
    first["path"].write_bytes(b"first package")
    first_record = store.record_package_version(
        project["id"],
        first["path"],
        version=first["version"],
        word_version={"version": 3, "filename": "测试项目_投标文件_V003.docx"},
        checklist_summary={
            "total": 6,
            "required": 5,
            "ready": 5,
            "not_applicable": 1,
            "pending_required": 0,
            "linked": 4,
            "broken_links": 0,
            "complete": True,
        },
        attachment_count=4,
        warning_count=1,
        internal_review_only=True,
        note="第一次内部预审",
    )
    second = store.next_package_version(project["id"], "测试项目_最终提交包.zip")
    second["path"].write_bytes(b"second package")
    store.record_package_version(
        project["id"],
        second["path"],
        version=second["version"],
        word_version={"version": 4, "filename": "测试项目_投标文件_V004.docx"},
        checklist_summary={"total": 6, "required": 5, "ready": 5, "complete": True},
        attachment_count=5,
    )

    versions = store.list_package_versions(project["id"])

    assert first["filename"].endswith("_P001.zip")
    assert second["filename"].endswith("_P002.zip")
    assert first_record["sha256"]
    assert [item["version"] for item in versions] == [2, 1]
    assert versions[1]["word_version"] == 3
    assert versions[1]["checklist_summary"]["required"] == 5
    assert versions[1]["attachment_count"] == 4
    assert versions[1]["warning_count"] == 1
    assert versions[1]["internal_review_only"] is True
    assert versions[1]["note"] == "第一次内部预审"


def test_submission_package_retention_keeps_latest_versions(tmp_path: Path) -> None:
    store = ProjectStore(tmp_path / "data")
    project = store.create_project("提交包版本上限测试")

    for index in range(MAX_PACKAGE_VERSIONS + 2):
        target = store.next_package_version(project["id"], "最终提交包.zip")
        target["path"].write_bytes(f"package-{index}".encode())
        store.record_package_version(
            project["id"],
            target["path"],
            version=target["version"],
            word_version={"version": 1, "filename": "投标文件_V001.docx"},
            checklist_summary={"total": 1, "required": 1, "ready": 1, "complete": True},
            attachment_count=0,
        )

    versions = store.list_package_versions(project["id"])
    package_files = list((store.project_dir(project["id"]) / "packages").glob("*.zip"))

    assert len(versions) == MAX_PACKAGE_VERSIONS
    assert len(package_files) == MAX_PACKAGE_VERSIONS
    assert versions[0]["version"] == MAX_PACKAGE_VERSIONS + 2
    assert all(item["version"] != 1 for item in versions)


def test_project_archive_round_trip_and_id_conflict(tmp_path: Path) -> None:
    source_store = ProjectStore(tmp_path / "source")
    source = source_store.create_project("迁移项目")
    source_store.save_source(source["id"], "招标文件.txt", "招标内容".encode())
    source_store.save_json(source["id"], "analysis", {"ok": True})
    source_store.save_json(source["id"], "analysis_baseline", {"baseline": True})
    source_store.save_json(
        source["id"],
        "analysis_baseline_meta",
        {"origin": "自动分析", "source_fingerprint": "abc"},
    )
    source_store.save_json(source["id"], "analysis_acceptance", {"complete": True})
    source_store.save_knowledge_file(source["id"], "company", "企业.txt", b"company")
    source_store.save_attachment_file(source["id"], "qualification", "营业执照.pdf", b"license")
    source_store.save_json(
        source["id"],
        "submission_checklist",
        [{"name": "营业执照", "required": True, "status": "已备妥"}],
    )
    output = source_store.output_path(source["id"], "初稿.docx")
    output.write_bytes(b"docx")
    source_store.record_export_version(
        source["id"],
        output,
        version=1,
        chapter_count=2,
        review_summary={"pending": 1, "high": 1, "medium": 0, "low": 0},
        note="备份版本",
    )
    package = source_store.package_path(source["id"], "提交包_P001.zip")
    package.write_bytes(b"package")
    source_store.record_package_version(
        source["id"],
        package,
        version=1,
        word_version={"version": 1, "filename": output.name},
        checklist_summary={"total": 1, "required": 1, "ready": 1, "complete": True},
        attachment_count=1,
        note="备份提交包",
    )
    backup = source_store.export_project_archive(source["id"])

    target_store = ProjectStore(tmp_path / "target")
    imported = target_store.import_project_archive(backup)

    assert imported["id"] == source["id"]
    assert imported["name"] == "迁移项目"
    assert imported["archived"] == 0
    assert target_store.load_json(imported["id"], "analysis") == {"ok": True}
    assert target_store.load_json(imported["id"], "analysis_baseline") == {"baseline": True}
    assert target_store.load_json(imported["id"], "analysis_baseline_meta") == {
        "origin": "自动分析",
        "source_fingerprint": "abc",
    }
    assert target_store.load_json(imported["id"], "analysis_acceptance") == {"complete": True}
    assert target_store.source_path(imported["id"]).read_text(encoding="utf-8") == "招标内容"
    assert target_store.output_path(imported["id"], "初稿.docx").read_bytes() == b"docx"
    assert target_store.list_export_versions(imported["id"])[0]["note"] == "备份版本"
    assert target_store.list_package_versions(imported["id"])[0]["note"] == "备份提交包"
    assert target_store.list_package_versions(imported["id"])[0]["path"].read_bytes() == b"package"
    assert target_store.list_attachment_files(imported["id"])["qualification"][0].read_bytes() == b"license"
    assert target_store.load_json(imported["id"], "submission_checklist")[0]["status"] == "已备妥"

    duplicate = target_store.import_project_archive(backup)
    assert duplicate["id"] != source["id"]
    assert duplicate["name"] == "迁移项目"


def test_project_archive_rejects_path_traversal(tmp_path: Path) -> None:
    manifest = {
        "format": "bid-assistant-project",
        "version": 1,
        "project": {"id": "proj_bad", "name": "坏备份"},
        "files": [],
    }
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("manifest.json", json.dumps(manifest))
        archive.writestr("project/../../outside.txt", "bad")

    store = ProjectStore(tmp_path / "data")
    with pytest.raises(ProjectArchiveError, match="不安全路径"):
        store.import_project_archive(buffer.getvalue())
    assert not (tmp_path / "outside.txt").exists()


@pytest.mark.parametrize("unsafe_path", ["project/C:/secret.txt", "project/CON.txt", "project/bad:name.txt"])
def test_project_archive_rejects_windows_unsafe_paths(tmp_path: Path, unsafe_path: str) -> None:
    manifest = {
        "format": "bid-assistant-project",
        "version": 1,
        "project": {"id": "proj_bad", "name": "坏备份"},
        "files": [],
    }
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("manifest.json", json.dumps(manifest))
        archive.writestr(unsafe_path, "bad")

    store = ProjectStore(tmp_path / "data")
    with pytest.raises(ProjectArchiveError, match="不安全路径"):
        store.import_project_archive(buffer.getvalue())


def test_project_archive_rejects_tampered_file(tmp_path: Path) -> None:
    manifest = {
        "format": "bid-assistant-project",
        "version": 1,
        "project": {"id": "proj_bad", "name": "坏备份"},
        "files": [{"path": "analysis.json", "size": 2, "sha256": "0" * 64}],
    }
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("manifest.json", json.dumps(manifest))
        archive.writestr("project/analysis.json", "{}")

    store = ProjectStore(tmp_path / "data")
    with pytest.raises(ProjectArchiveError, match="完整性校验失败"):
        store.import_project_archive(buffer.getvalue())
