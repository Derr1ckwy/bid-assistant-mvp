from __future__ import annotations

import argparse
import io
import json

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from bid_assistant.docx_quality import verify_docx_output
from bid_assistant.exporter import export_docx
from bid_assistant.models import ChapterDraft, ReviewReport, TenderAnalysis
from bid_assistant.packager import (
    build_package_readiness,
    create_submission_package,
    verify_submission_package,
)
from bid_assistant.storage import ProjectStore
from bid_assistant.submission import (
    ATTACHMENT_CATEGORY_LABELS,
    build_attachment_inventory,
    summarize_submission_items,
)


NOTICE = "模拟测试资料，不可用于真实投标。"


def _make_docx_bytes(title: str, lines: list[str]) -> bytes:
    document = Document()
    for style_name in ("Normal", "Title"):
        style = document.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    document.styles["Normal"].font.size = Pt(12)
    document.add_heading(title, level=0)
    document.add_paragraph(NOTICE)
    for line in lines:
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _unique_project_name(store: ProjectStore, requested: str) -> str:
    existing_names = {item["name"] for item in store.list_projects(include_archived=True)}
    if requested not in existing_names:
        return requested
    sequence = 2
    while f"{requested}-{sequence}" in existing_names:
        sequence += 1
    return f"{requested}-{sequence}"


def run_simulation(source_project_id: str, requested_name: str) -> dict:
    store = ProjectStore("data")
    project_name = _unique_project_name(store, requested_name)
    project = store.duplicate_project(source_project_id, name=project_name)
    project_id = project["id"]

    analysis = TenderAnalysis.model_validate(store.load_json(project_id, "analysis"))
    drafts = [
        ChapterDraft.model_validate(item)
        for item in store.load_json(project_id, "drafts", [])
    ]
    review = ReviewReport.model_validate(
        store.load_json(project_id, "review", {"issues": []})
    )

    target = store.next_output_version(project_id, f"{project_name}_投标文件初稿.docx")
    export_docx(
        target["path"],
        analysis,
        drafts,
        review,
        qualification_images=store.list_qualification_images(project_id),
        include_internal_appendices=False,
    )
    review_summary = {
        "pending": review.pending_count(),
        "high": review.severity_count("高"),
        "medium": review.severity_count("中"),
        "low": review.severity_count("低"),
    }
    word_version = store.record_export_version(
        project_id,
        target["path"],
        version=target["version"],
        chapter_count=len(drafts),
        review_summary=review_summary,
        warning_count=1 if review.pending_count() else 0,
        note="交付打包流程模拟测试；仅供内部预审",
        qualification_image_count=len(store.list_qualification_images(project_id)),
    )
    word_quality = verify_docx_output(
        word_version["path"],
        expected_sha256=word_version["sha256"],
        expected_chapter_count=len(drafts),
        template_mode=False,
    )
    word_only_readiness = build_package_readiness(
        word_version,
        [],
        set(),
        review,
        word_quality=word_quality,
    )

    attachment_specs = [
        (
            "qualification",
            "模拟测试_不可用于真实投标_营业执照说明.docx",
            "营业执照复印件",
            ["本文件仅用于验证资格附件上传、关联和打包流程。"],
        ),
        (
            "qualification",
            "模拟测试_不可用于真实投标_建筑业企业资质说明.docx",
            "建筑业企业资质证明",
            ["资质名称、等级和有效期均须在真实投标前以原件核验。"],
        ),
        (
            "business",
            "模拟测试_不可用于真实投标_法定代表人授权委托书.docx",
            "法定代表人授权委托书",
            ["授权人、被授权人、身份证号、日期及签章位置均为测试占位。"],
        ),
        (
            "technical",
            "模拟测试_不可用于真实投标_技术响应说明.docx",
            "技术响应文件",
            ["技术响应正文由系统生成 Word 承载，本附件仅用于演示补充附件。"],
        ),
        (
            "signature",
            "模拟测试_不可用于真实投标_签字盖章检查表.docx",
            "签字盖章及装订检查",
            ["真实提交前必须人工核对签字、盖章、骑缝章、密封和装订要求。"],
        ),
    ]
    for category, filename, title, lines in attachment_specs:
        store.save_attachment_file(
            project_id,
            category,
            filename,
            _make_docx_bytes(title, lines),
        )

    store.save_attachment_file(
        project_id,
        "pricing",
        "模拟测试_不可用于真实投标_报价响应表.csv",
        (
            "\ufeff项目,单位,数量,单价,合价,说明\r\n"
            "模拟测试项,项,1,,,不可用于真实投标\r\n"
        ).encode("utf-8"),
    )
    attachment_files = store.list_attachment_files(project_id)
    attachment_refs = {
        f"{ATTACHMENT_CATEGORY_LABELS[category_id]}/{path.name}"
        for category_id, paths in attachment_files.items()
        for path in paths
    }
    inventory_items = build_attachment_inventory(attachment_files)
    inventory_summary = summarize_submission_items(inventory_items, attachment_refs)
    attachment_readiness = build_package_readiness(
        word_version,
        inventory_items,
        attachment_refs,
        review,
        word_quality=word_quality,
    )
    if not attachment_readiness["can_package"]:
        raise RuntimeError(f"模拟交付包仍不可生成：{attachment_readiness}")

    package_target = store.next_package_version(
        project_id,
        f"{project_name}_交付包.zip",
    )
    create_submission_package(
        package_target["path"],
        project=project,
        word_version=word_version,
        items=inventory_items,
        attachment_files=attachment_files,
        review_summary=review_summary,
        internal_review_only=True,
        note="模拟测试资料，不可用于真实投标；仅验证交付打包流程",
    )
    package_version = store.record_package_version(
        project_id,
        package_target["path"],
        version=package_target["version"],
        word_version=word_version,
        checklist_summary=inventory_summary,
        attachment_count=sum(len(paths) for paths in attachment_files.values()),
        warning_count=attachment_readiness["warning_count"],
        internal_review_only=True,
        note="模拟测试资料，不可用于真实投标；仅验证交付打包流程",
    )
    verification = verify_submission_package(
        package_version["path"],
        expected_sha256=package_version["sha256"],
    )
    if not verification["valid"]:
        raise RuntimeError(f"模拟提交包校验失败：{verification['errors']}")
    store.update_project(project_id, status="packaged")

    report = {
        "project_id": project_id,
        "project_name": project_name,
        "source_project_id": source_project_id,
        "rag_file_count": sum(
            len(paths) for paths in store.list_knowledge_files(project_id).values()
        ),
        "word": {
            "filename": word_version["filename"],
            "version": word_version["version"],
            "quality_valid": word_quality["valid"],
            "quality_warnings": word_quality["warnings"],
        },
        "stages": {
            "word_only": {
                "can_package": word_only_readiness["can_package"],
                "requires_confirmation": word_only_readiness["requires_confirmation"],
            },
            "with_attachments": {
                "can_package": attachment_readiness["can_package"],
                "requires_confirmation": attachment_readiness["requires_confirmation"],
                "warning_count": attachment_readiness["warning_count"],
            },
        },
        "file_inventory": inventory_summary,
        "review": review_summary,
        "package": {
            "filename": package_version["filename"],
            "version": package_version["version"],
            "internal_review_only": package_version["internal_review_only"],
            "verification_valid": verification["valid"],
            "file_count": verification["file_count"],
            "sha256": verification["sha256"],
        },
    }
    store.save_json(project_id, "submission_simulation_report", report)

    walkthrough = f"""# 交付打包模拟测试说明

> {NOTICE}

## 本次模拟结果

- 项目：{project_name}
- 项目 ID：`{project_id}`
- 继承 RAG 知识文件：{report['rag_file_count']} 份
- Word：V{word_version['version']:03d}，成品质检 {'通过' if word_quality['valid'] else '未通过'}
- 模拟补充附件：{sum(len(paths) for paths in attachment_files.values())} 个
- 系统自动文件目录：{inventory_summary['total']} 项
- 交付包：P{package_version['version']:03d}，完整性校验 {'通过' if verification['valid'] else '未通过'}
- 用途：仅内部预审，不可用于真实投标

## 你在页面中的正确操作顺序

1. 打开“7. 交付打包”。
2. 如有需要，上传营业执照、资质证书、授权书、报价表等补充附件；没有附件时可以跳过。
3. 选择一个 Word 版本。
4. 确认本次包内容；系统会自动生成文件目录，无需维护材料清单。
5. 若存在复核风险，确认本包仅用于内部预审，然后生成交付包。

## 三类资料不要混淆

- RAG 知识资料：用于生成正文和检索事实，不会自动进入交付包。
- 补充附件：需要与 Word 一起交付的营业执照、授权书、报价表、证书等文件。
- 自动文件目录：系统在打包时生成，用户不需要编辑或上传。
"""
    (store.project_dir(project_id) / "交付打包模拟测试说明.md").write_text(
        walkthrough,
        encoding="utf-8",
    )
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="创建隔离的交付打包流程模拟项目")
    parser.add_argument("source_project_id", help="作为业务数据来源的项目 ID")
    parser.add_argument(
        "--name",
        default="交付打包模拟测试",
        help="模拟项目名称；重名时自动追加序号",
    )
    args = parser.parse_args()
    report = run_simulation(args.source_project_id, args.name)
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
