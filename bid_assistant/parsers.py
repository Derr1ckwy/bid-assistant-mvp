from __future__ import annotations

import csv
import io
import json
import re
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from pypdf import PdfReader

from bid_assistant.models import ParsedDocument, ParsedPage


TENDER_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
KNOWLEDGE_EXTENSIONS = TENDER_EXTENSIONS | {".csv", ".json"}
SUPPORTED_EXTENSIONS = KNOWLEDGE_EXTENSIONS

MAX_JSON_VALUES = 50000
MAX_JSON_DEPTH = 40


class DocumentParseError(RuntimeError):
    pass


def _decode_text(path: Path) -> str:
    content = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _parse_pdf(path: Path) -> tuple[list[ParsedPage], list[str]]:
    warnings: list[str] = []
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise DocumentParseError(f"PDF 无法读取：{exc}") from exc

    pages = [
        ParsedPage(page_number=index, text=(page.extract_text() or "").strip())
        for index, page in enumerate(reader.pages, start=1)
    ]
    if reader.is_encrypted:
        warnings.append("PDF 带有加密标记，部分内容可能无法解析。")
    return pages, warnings


def _parse_docx(path: Path) -> list[ParsedPage]:
    try:
        document = Document(str(path))
    except Exception as exc:
        raise DocumentParseError(f"DOCX 无法读取：{exc}") from exc

    blocks: list[str] = []
    blocks.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return [ParsedPage(page_number=None, text="\n".join(blocks))]


def _unique_headers(values: list[str]) -> list[str]:
    result: list[str] = []
    counts: dict[str, int] = {}
    for index, raw_value in enumerate(values, start=1):
        value = re.sub(r"\s+", " ", raw_value).strip() or f"字段{index}"
        counts[value] = counts.get(value, 0) + 1
        suffix = f" ({counts[value]})" if counts[value] > 1 else ""
        result.append(f"{value}{suffix}")
    return result


def _parse_csv(path: Path) -> tuple[list[ParsedPage], list[str]]:
    text = _decode_text(path)
    if not text.strip():
        return [ParsedPage(page_number=1, text="")], ["CSV 文件为空。"]

    sample = text[:8192]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
    except csv.Error:
        dialect = csv.excel

    try:
        reader = csv.reader(io.StringIO(text, newline=""), dialect=dialect, strict=True)
        rows = [
            [re.sub(r"\s+", " ", cell).strip() for cell in row]
            for row in reader
            if any(cell.strip() for cell in row)
        ]
    except csv.Error as exc:
        raise DocumentParseError(f"CSV 无法读取：{exc}") from exc

    if not rows:
        return [ParsedPage(page_number=1, text="")], ["CSV 没有可用数据行。"]

    headers = _unique_headers(rows[0])
    blocks = ["CSV 字段：" + " | ".join(headers)]
    for row_index, row in enumerate(rows[1:], start=1):
        padded = row + [""] * max(0, len(headers) - len(row))
        pairs = [
            f"{headers[index]}: {value}"
            for index, value in enumerate(padded[: len(headers)])
            if value
        ]
        if len(row) > len(headers):
            pairs.extend(
                f"字段{index + 1}: {value}"
                for index, value in enumerate(row[len(headers) :], start=len(headers))
                if value
            )
        if pairs:
            blocks.append(f"记录 {row_index}\n" + "\n".join(pairs))

    return [ParsedPage(page_number=1, text="\n\n".join(blocks))], []


def _json_scalar(value: object) -> str:
    if isinstance(value, str):
        return re.sub(r"\s+", " ", value).strip()
    return json.dumps(value, ensure_ascii=False)


def _flatten_json(
    value: object,
    path: str,
    lines: list[str],
    *,
    depth: int = 0,
) -> bool:
    if len(lines) >= MAX_JSON_VALUES:
        return True
    if depth > MAX_JSON_DEPTH:
        lines.append(f"{path}: [嵌套层级超过 {MAX_JSON_DEPTH}，已停止展开]")
        return False

    if isinstance(value, dict):
        if not value:
            lines.append(f"{path}: {{}}")
            return False
        for key, child in value.items():
            child_path = f"{path}.{key}" if path else str(key)
            if _flatten_json(child, child_path, lines, depth=depth + 1):
                return True
        return False

    if isinstance(value, list):
        if not value:
            lines.append(f"{path}: []")
            return False
        for index, child in enumerate(value):
            child_path = f"{path}[{index}]" if path else f"记录[{index}]"
            if _flatten_json(child, child_path, lines, depth=depth + 1):
                return True
        return False

    lines.append(f"{path or '值'}: {_json_scalar(value)}")
    return False


def _parse_json(path: Path) -> tuple[list[ParsedPage], list[str]]:
    text = _decode_text(path)
    try:
        payload = json.loads(text)
    except json.JSONDecodeError as exc:
        raise DocumentParseError(
            f"JSON 无法读取：第 {exc.lineno} 行、第 {exc.colno} 列格式错误。"
        ) from exc

    lines: list[str] = []
    truncated = _flatten_json(payload, "", lines)
    warnings: list[str] = []
    if payload == {} or payload == []:
        warnings.append("JSON 为空对象或空数组。")
    if truncated:
        warnings.append(f"JSON 字段超过 {MAX_JSON_VALUES} 个，仅保留前 {MAX_JSON_VALUES} 个值。")
    return [ParsedPage(page_number=1, text="\n".join(lines))], warnings


def _parse_document_native(path: str | Path) -> ParsedDocument:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(f"暂不支持 {suffix or '未知'} 格式。")

    warnings: list[str] = []
    if suffix == ".pdf":
        pages, pdf_warnings = _parse_pdf(source)
        warnings.extend(pdf_warnings)
    elif suffix == ".docx":
        pages = _parse_docx(source)
    elif suffix == ".csv":
        pages, csv_warnings = _parse_csv(source)
        warnings.extend(csv_warnings)
    elif suffix == ".json":
        pages, json_warnings = _parse_json(source)
        warnings.extend(json_warnings)
    else:
        pages = [ParsedPage(page_number=1, text=_decode_text(source).strip())]

    full_text = "\n\n".join(
        f"[第 {page.page_number} 页]\n{page.text}" if page.page_number else page.text
        for page in pages
        if page.text
    ).strip()
    char_count = sum(len(page.text) for page in pages)
    possible_scanned = suffix == ".pdf" and char_count < max(80, len(pages) * 30)
    if possible_scanned:
        warnings.append("页面文本过少，可能是扫描件。请接入 MinerU/RAGFlow OCR 后再分析。")
    if not full_text:
        warnings.append("没有提取到可分析文本。")

    return ParsedDocument(
        filename=source.name,
        file_type=suffix.lstrip("."),
        pages=pages,
        full_text=full_text,
        char_count=char_count,
        possible_scanned_document=possible_scanned,
        parser_engine="native",
        warnings=warnings,
    )


def parse_document(path: str | Path, *, mode: str = "native", mineru_client=None) -> ParsedDocument:
    if mode not in {"native", "auto", "mineru"}:
        raise ValueError("Invalid parser mode")
    native = _parse_document_native(path)
    should_try_mineru = mode == "mineru" or (
        mode == "auto" and native.file_type == "pdf" and native.possible_scanned_document
    )
    if not should_try_mineru:
        return native
    if mineru_client is None:
        from bid_assistant.config import settings
        from bid_assistant.ocr import MinerUClient

        mineru_client = MinerUClient(settings)
    try:
        enhanced = mineru_client.parse(path)
        if enhanced.char_count <= native.char_count and native.full_text.strip():
            native.warnings.append("MinerU 未提取出更多有效文本，本次保留原生解析结果。")
            return native
        return enhanced
    except Exception as exc:
        native.warnings.append(f"MinerU 增强解析不可用，已保留原生结果：{exc}")
        return native


def parse_document_bytes(filename: str, content: bytes, *, mode: str = "native", mineru_client=None) -> ParsedDocument:
    safe_name = Path(filename).name
    if not safe_name:
        raise DocumentParseError("文件名为空。")
    with TemporaryDirectory(prefix="bid-assistant-parse-") as temporary_dir:
        path = Path(temporary_dir) / safe_name
        path.write_bytes(content)
        return parse_document(path, mode=mode, mineru_client=mineru_client)
