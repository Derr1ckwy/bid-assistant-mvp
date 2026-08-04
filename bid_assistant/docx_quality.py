from __future__ import annotations

import hashlib
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm


_SHA256_PATTERN = re.compile(r"[0-9a-f]{64}")
_CHAPTER_HEADING_PATTERN = re.compile(r"^第\s*\d+\s*章")
_PLACEHOLDER_PATTERN = re.compile(r"待补充|待确认(?!项)|待核对(?!事项)|【[^】]{0,80}】")
_REQUIRED_PARTS = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}


def _sha256_path(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        while chunk := source.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def _document_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(value for value in values if value)


def _is_close(value, expected, tolerance) -> bool:
    return value is not None and abs(value - expected) <= tolerance


def verify_docx_output(
    docx_path: str | Path,
    *,
    expected_sha256: str = "",
    expected_chapter_count: int = 0,
) -> dict:
    target = Path(docx_path)
    result = {
        "valid": False,
        "verified_at": datetime.now(timezone.utc).isoformat(timespec="milliseconds"),
        "filename": target.name,
        "size": 0,
        "sha256": "",
        "expected_sha256": expected_sha256,
        "paragraph_count": 0,
        "table_count": 0,
        "section_count": 0,
        "chapter_count": 0,
        "placeholder_count": 0,
        "errors": [],
        "warnings": [],
        "checks": [],
    }
    errors: list[str] = result["errors"]
    warnings: list[str] = result["warnings"]
    checks: list[dict] = result["checks"]

    if not target.is_file():
        errors.append("Word 文件不存在。")
        return result
    if target.suffix.lower() != ".docx":
        errors.append("文件扩展名不是 DOCX。")
        return result

    try:
        result["size"] = target.stat().st_size
        result["sha256"] = _sha256_path(target)
    except OSError:
        errors.append("Word 文件无法读取。")
        return result

    if expected_sha256:
        if not _SHA256_PATTERN.fullmatch(expected_sha256):
            errors.append("版本记录中的 Word SHA-256 格式无效。")
        elif result["sha256"] != expected_sha256:
            errors.append("Word 文件 SHA-256 与版本记录不一致。")
    else:
        warnings.append("Word 版本记录缺少 SHA-256，无法确认文件是否在导出后被替换。")

    try:
        with zipfile.ZipFile(target) as archive:
            parts = {info.filename for info in archive.infolist() if not info.is_dir()}
    except (OSError, zipfile.BadZipFile):
        errors.append("文件不是有效的 DOCX 压缩包。")
        return result

    missing_parts = sorted(_REQUIRED_PARTS - parts)
    if missing_parts:
        errors.append(f"DOCX 缺少必要结构：{'、'.join(missing_parts)}")
        return result

    try:
        document = Document(target)
    except Exception:
        errors.append("python-docx 无法打开该文件，文档结构可能已损坏。")
        return result

    result["paragraph_count"] = len(document.paragraphs)
    result["table_count"] = len(document.tables)
    result["section_count"] = len(document.sections)
    text = _document_text(document)
    if not text.strip():
        errors.append("Word 文件没有可识别的正文内容。")

    chapter_count = sum(
        paragraph.style is not None
        and paragraph.style.name == "Heading 1"
        and bool(_CHAPTER_HEADING_PATTERN.match(paragraph.text.strip()))
        for paragraph in document.paragraphs
    )
    result["chapter_count"] = chapter_count
    if expected_chapter_count > 0 and chapter_count < expected_chapter_count:
        errors.append(
            f"Word 仅识别到 {chapter_count} 个正文章节，少于版本记录的 {expected_chapter_count} 个。"
        )

    a4_sections = 0
    margin_warnings = 0
    for section in document.sections:
        is_a4 = _is_close(section.page_width, Cm(21), Cm(0.2)) and _is_close(
            section.page_height, Cm(29.7), Cm(0.2)
        )
        a4_sections += int(is_a4)
        expected_margins = (
            (section.top_margin, Cm(2.5)),
            (section.bottom_margin, Cm(2.5)),
            (section.left_margin, Cm(2.8)),
            (section.right_margin, Cm(2.5)),
        )
        if not all(_is_close(value, expected, Cm(0.25)) for value, expected in expected_margins):
            margin_warnings += 1
    if document.sections and a4_sections != len(document.sections):
        errors.append(f"有 {len(document.sections) - a4_sections} 个分节不是 A4 纵向页面。")
    checks.append(
        {
            "label": "页面规格",
            "status": "pass" if document.sections and a4_sections == len(document.sections) else "block",
            "detail": f"A4 分节 {a4_sections}/{len(document.sections)}。",
        }
    )
    if margin_warnings:
        warnings.append(f"有 {margin_warnings} 个分节的页边距偏离系统标准版式。")

    try:
        normal = document.styles["Normal"]
    except KeyError:
        normal = None
    normal_font_ok = bool(
        normal is not None
        and normal.font.size is not None
        and abs(normal.font.size.pt - 10.5) <= 0.2
    )
    normal_spacing_ok = bool(
        normal is not None and normal.paragraph_format.line_spacing == 1.5
    )
    normal_east_asia_ok = bool(
        normal is not None and 'w:eastAsia="宋体"' in normal._element.xml
    )
    if not normal_font_ok or not normal_east_asia_ok:
        warnings.append("正文样式不是宋体 10.5 磅。")
    if not normal_spacing_ok:
        warnings.append("正文样式不是 1.5 倍行距。")
    checks.append(
        {
            "label": "正文样式",
            "status": "pass" if normal_font_ok and normal_spacing_ok and normal_east_asia_ok else "warning",
            "detail": "宋体 10.5 磅、1.5 倍行距。",
        }
    )

    missing_spacing_styles = []
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = document.styles[style_name]
        except KeyError:
            missing_spacing_styles.append(style_name)
            continue
        r_pr = style._element.get_or_add_rPr()
        if r_pr.find(qn("w:spacing")) is None:
            missing_spacing_styles.append(style_name)
    if missing_spacing_styles:
        warnings.append(f"标题样式缺少明确字符间距：{'、'.join(missing_spacing_styles)}。")
    checks.append(
        {
            "label": "标题字距",
            "status": "warning" if missing_spacing_styles else "pass",
            "detail": "标题样式包含明确字符间距。",
        }
    )

    page_number_ready = any(" PAGE " in section.footer._element.xml for section in document.sections)
    first_page_ready = all(
        section.different_first_page_header_footer for section in document.sections
    ) if document.sections else False
    if not page_number_ready:
        warnings.append("页脚未识别到 PAGE 页码字段。")
    if not first_page_ready:
        warnings.append("文档未对封面启用首页不同。")
    checks.append(
        {
            "label": "页眉页脚",
            "status": "pass" if page_number_ready and first_page_ready else "warning",
            "detail": "包含页码字段并启用首页不同。",
        }
    )

    fixed_table_count = 0
    no_split_table_count = 0
    split_check_tables = []
    for table in document.tables:
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        width = table._tbl.tblPr.find(qn("w:tblW"))
        if (
            layout is not None
            and layout.get(qn("w:type")) == "fixed"
            and width is not None
            and width.get(qn("w:type")) == "dxa"
        ):
            fixed_table_count += 1
        if table.style is not None and table.style.name == "Table Grid":
            split_check_tables.append(table)
            if all(
                row._tr.trPr is not None and row._tr.trPr.find(qn("w:cantSplit")) is not None
                for row in table.rows
            ):
                no_split_table_count += 1
    if document.tables and fixed_table_count != len(document.tables):
        warnings.append(f"有 {len(document.tables) - fixed_table_count} 个表格未使用固定宽度。")
    if split_check_tables and no_split_table_count != len(split_check_tables):
        warnings.append(f"有 {len(split_check_tables) - no_split_table_count} 个正文表格允许单行跨页拆分。")
    checks.append(
        {
            "label": "表格版式",
            "status": (
                "pass"
                if not document.tables
                or (
                    fixed_table_count == len(document.tables)
                    and no_split_table_count == len(split_check_tables)
                )
                else "warning"
            ),
            "detail": (
                f"固定宽度表格 {fixed_table_count}/{len(document.tables)}，"
                f"禁止拆行正文表格 {no_split_table_count}/{len(split_check_tables)}。"
            ),
        }
    )

    if "投 标 响 应 文 件" not in text:
        warnings.append("未识别到系统标准封面标题。")
    if "目 录" not in text:
        warnings.append("未识别到目录标题。")
    update_fields_ready = document.settings._element.find(qn("w:updateFields")) is not None
    if not update_fields_ready:
        warnings.append("文档未启用打开时更新字段。")

    placeholders = _PLACEHOLDER_PATTERN.findall(text)
    result["placeholder_count"] = len(placeholders)
    if placeholders:
        warnings.append(f"Word 中仍有 {len(placeholders)} 处待补充、待确认或占位内容。")

    checks.append(
        {
            "label": "内容占位符",
            "status": "warning" if placeholders else "pass",
            "detail": f"识别到 {len(placeholders)} 处占位内容。",
        }
    )
    result["valid"] = not errors
    return result


def build_docx_quality_report(
    quality: dict,
    *,
    word_version: dict | None = None,
) -> bytes:
    conclusion = "通过" if quality.get("valid") else "不通过"
    lines = [
        "Word 成品完整性与版式结构质检报告",
        "=" * 36,
        f"质检结论：{conclusion}",
        f"质检时间：{quality.get('verified_at', '-')}",
        f"文件名称：{quality.get('filename', '-')}",
        f"文件大小：{int(quality.get('size', 0))} 字节",
        f"实际 SHA-256：{quality.get('sha256', '-')}",
        f"记录 SHA-256：{quality.get('expected_sha256') or '-'}",
        f"段落数：{int(quality.get('paragraph_count', 0))}",
        f"表格数：{int(quality.get('table_count', 0))}",
        f"分节数：{int(quality.get('section_count', 0))}",
        f"正文章节数：{int(quality.get('chapter_count', 0))}",
        f"占位内容数：{int(quality.get('placeholder_count', 0))}",
    ]
    if word_version:
        lines.extend(
            [
                f"Word 版本：V{int(word_version.get('version', 0)):03d}",
                f"版本说明：{word_version.get('note') or '-'}",
            ]
        )

    lines.extend(["", "结构检查："])
    for item in quality.get("checks") or []:
        status_text = {"pass": "通过", "warning": "需确认", "block": "不通过"}.get(
            item.get("status"), "未知"
        )
        lines.append(f"- [{status_text}] {item.get('label', '-')}：{item.get('detail', '-')}")
    if not quality.get("checks"):
        lines.append("- 无可用检查结果")

    lines.extend(["", "错误项："])
    errors = quality.get("errors") or []
    lines.extend(f"- {item}" for item in errors)
    if not errors:
        lines.append("- 无")

    lines.extend(["", "提示项："])
    warnings = quality.get("warnings") or []
    lines.extend(f"- {item}" for item in warnings)
    if not warnings:
        lines.append("- 无")

    lines.extend(
        [
            "",
            "说明：本报告检查文件完整性和系统标准版式结构，不替代人工逐页审阅及甲方模板核对。",
        ]
    )
    return ("\ufeff" + "\r\n".join(lines) + "\r\n").encode("utf-8")
