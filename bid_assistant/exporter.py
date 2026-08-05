from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

from bid_assistant.models import ChapterDraft, ReviewReport, TenderAnalysis


BODY_FONT = "Times New Roman"
BODY_FONT_EAST_ASIA = "宋体"
HEADING_FONT = BODY_FONT
HEADING_FONT_EAST_ASIA = BODY_FONT_EAST_ASIA
COVER_FONT = BODY_FONT
COVER_FONT_EAST_ASIA = BODY_FONT_EAST_ASIA
CONTENT_WIDTH_DXA = 8901
TABLE_INDENT_DXA = 0
BLACK = "000000"
DARK_TEXT = "1A1A1A"
MUTED_TEXT = "666666"
LIGHT_GRAY = "F2F2F2"
PALE_GRAY = "FAFAFA"
BORDER = "7F7F7F"


def _set_rpr_font_identity(r_pr, *, latin: str, east_asia: str) -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        r_fonts.attrib.pop(qn(f"w:{key}"), None)
    for key, value in (("ascii", latin), ("hAnsi", latin), ("cs", latin), ("eastAsia", east_asia)):
        r_fonts.set(qn(f"w:{key}"), value)
    r_fonts.set(qn("w:hint"), "eastAsia")

    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    for key in ("val", "eastAsia", "bidi"):
        language.set(qn(f"w:{key}"), "zh-CN")


def _set_character_spacing(element, points: float) -> None:
    r_pr = element.get_or_add_rPr()
    for existing in r_pr.findall(qn("w:spacing")):
        r_pr.remove(existing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(round(points * 20)))
    r_pr.append(spacing)


def _set_style_font(
    style,
    *,
    latin: str,
    east_asia: str,
    size: float,
    bold: bool = False,
    color: str = DARK_TEXT,
    character_spacing: float = 0,
) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style._element.get_or_add_rPr()
    _set_rpr_font_identity(r_pr, latin=latin, east_asia=east_asia)
    _set_character_spacing(style._element, character_spacing)


def _format_run(
    run,
    *,
    latin: str = BODY_FONT,
    east_asia: str = BODY_FONT_EAST_ASIA,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    character_spacing: float = 0,
) -> None:
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    _set_rpr_font_identity(r_pr, latin=latin, east_asia=east_asia)
    _set_character_spacing(run._element, character_spacing)


def _set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for existing in p_pr.findall(qn("w:shd")):
        p_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _set_paragraph_border(paragraph, *, side: str, color: str, size: int = 12, space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles_root = document.styles.element
    doc_defaults = styles_root.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_root.insert(0, doc_defaults)
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.insert(0, r_pr_default)
    default_r_pr = r_pr_default.find(qn("w:rPr"))
    if default_r_pr is None:
        default_r_pr = OxmlElement("w:rPr")
        r_pr_default.append(default_r_pr)
    _set_rpr_font_identity(default_r_pr, latin=BODY_FONT, east_asia=BODY_FONT_EAST_ASIA)

    normal = document.styles["Normal"]
    _set_style_font(
        normal,
        latin=BODY_FONT,
        east_asia=BODY_FONT_EAST_ASIA,
        size=12,
        color=DARK_TEXT,
    )
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = Pt(26)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.85)
    normal.paragraph_format.widow_control = True

    style_tokens = (
        ("Title", 26, 0, 12),
        ("Subtitle", 16, 0, 14),
        ("Heading 1", 18, 18, 12),
        ("Heading 2", 15, 14, 8),
        ("Heading 3", 13, 10, 6),
    )
    for name, size, before, after in style_tokens:
        style = document.styles[name]
        _set_style_font(
            style,
            latin=HEADING_FONT,
            east_asia=HEADING_FONT_EAST_ASIA,
            size=size,
            bold=True,
            color=BLACK,
            character_spacing=0,
        )
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        _set_style_font(
            style,
            latin=BODY_FONT,
            east_asia=BODY_FONT_EAST_ASIA,
            size=12,
            color=DARK_TEXT,
        )
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = Pt(24)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.widow_control = True

    settings = document.settings._element
    if settings.find(qn("w:updateFields")) is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def _create_numbering(document: Document, *, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "600")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "600")
    indent.set(qn("w:hanging"), "300")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([level, num])
    p_pr.append(num_pr)


def _add_page_number(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = None
    prefix = paragraph.add_run("第 ")
    _format_run(prefix, size=9, color=MUTED_TEXT)
    run = paragraph.add_run()
    _format_run(run, size=9, color=MUTED_TEXT)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])
    suffix = paragraph.add_run(" 页")
    _format_run(suffix, size=9, color=MUTED_TEXT)


def _add_header_footer(document: Document, project_name: str) -> None:
    for section in document.sections:
        section.different_first_page_header_footer = True
        first_header = section.first_page_header
        first_header.paragraphs[0].text = ""
        first_footer = section.first_page_footer
        first_footer.paragraphs[0].text = ""

        header = section.header
        header_table = header.add_table(rows=1, cols=2, width=Cm(15.7))
        _set_table_geometry(header_table, [6900, 2001], indent=0)
        _remove_table_borders(header_table)
        _set_table_border_side(header_table, side="bottom", color="BFBFBF", size=4)
        left_header, right_header = header_table.rows[0].cells
        for cell in (left_header, right_header):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=0, bottom=55, start=0, end=0)

        left_paragraph = left_header.paragraphs[0]
        left_paragraph.paragraph_format.first_line_indent = None
        left_paragraph.paragraph_format.space_after = Pt(0)
        left = left_paragraph.add_run(project_name)
        _format_run(left, size=8.5, color=MUTED_TEXT)

        right_paragraph = right_header.paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_paragraph.paragraph_format.first_line_indent = None
        right_paragraph.paragraph_format.space_after = Pt(0)
        right = right_paragraph.add_run("投标文件")
        _format_run(right, size=8.5, color=MUTED_TEXT)
        header._element.remove(header.paragraphs[0]._p)

        footer = section.footer
        table = footer.add_table(rows=1, cols=3, width=Cm(15.7))
        _set_table_geometry(table, [3000, 2901, 3000], indent=0)
        _remove_table_borders(table)
        left_cell, center_cell, right_cell = table.rows[0].cells
        for cell in (left_cell, center_cell, right_cell):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=40, bottom=0, start=0, end=0)

        page_paragraph = center_cell.paragraphs[0]
        page_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_paragraph.paragraph_format.space_before = Pt(2)
        page_paragraph.paragraph_format.space_after = Pt(0)
        _add_page_number(page_paragraph)
        footer._element.remove(footer.paragraphs[0]._p)


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_table_border_side(table, *, side: str, color: str, size: int) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    border = borders.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)


def _set_cell_margins(cell, *, top: int = 90, bottom: int = 90, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def _set_table_geometry(table, widths: list[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    total_width = sum(widths)
    if total_width <= 0 or len(widths) != len(table.columns):
        raise ValueError("Table widths must be positive and match the column count")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, attributes in (
        ("w:tblW", {"w:w": str(total_width), "w:type": "dxa"}),
        ("w:tblInd", {"w:w": str(indent), "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        existing = tbl_pr.find(qn(tag))
        if existing is None:
            existing = OxmlElement(tag)
            tbl_pr.append(existing)
        for key, value in attributes.items():
            existing.set(qn(key), value)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def _style_table(
    table,
    widths: list[int],
    *,
    center_columns: set[int] | None = None,
    nowrap_columns: set[int] | None = None,
    compact: bool = False,
    shade_alternate_rows: bool = False,
) -> None:
    center_columns = center_columns or set()
    nowrap_columns = nowrap_columns or set()
    table.style = "Table Grid"
    _set_table_geometry(table, widths)
    _set_table_borders(table)
    _repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        _prevent_row_split(row)
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if column_index in nowrap_columns:
                _set_cell_no_wrap(cell)
            _set_cell_margins(
                cell,
                top=50 if compact else 110,
                bottom=50 if compact else 110,
                start=90 if compact else 140,
                end=90 if compact else 140,
            )
            if row_index == 0:
                _set_cell_shading(cell, LIGHT_GRAY)
                if len(cell.text.strip()) <= 6:
                    _set_cell_no_wrap(cell)
            elif shade_alternate_rows and row_index % 2 == 0:
                _set_cell_shading(cell, PALE_GRAY)
            else:
                _set_cell_shading(cell, "FFFFFF")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1 if compact else 1.2
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if row_index == 0 or column_index in center_columns
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    _format_run(
                        run,
                        size=12,
                        bold=True if row_index == 0 else None,
                        color=BLACK,
                    )


def _column_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    column_count = len(headers)
    if column_count == 1:
        return [CONTENT_WIDTH_DXA]
    narrow_headers = {"序号", "编号", "页码", "分值", "得分", "数量", "状态", "级别"}
    minimums = [700 if header.strip() in narrow_headers else 1000 for header in headers]
    if sum(minimums) >= CONTENT_WIDTH_DXA:
        minimums = [CONTENT_WIDTH_DXA // column_count] * column_count
    text_lengths: list[int] = []
    for index, header in enumerate(headers):
        values = [header, *[row[index] if index < len(row) else "" for row in rows]]
        if header.strip() in narrow_headers:
            text_lengths.append(1)
        else:
            text_lengths.append(max(4, min(48, max(len(value) for value in values))))
    remaining = CONTENT_WIDTH_DXA - sum(minimums)
    weight_total = sum(text_lengths)
    widths = [minimum + int(remaining * weight / weight_total) for minimum, weight in zip(minimums, text_lengths, strict=True)]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _center_columns(headers: list[str]) -> set[int]:
    centered = {"序号", "编号", "页码", "分值", "得分", "数量", "状态", "级别", "是否响应"}
    return {index for index, header in enumerate(headers) if header.strip() in centered}


def _add_cover(document: Document, analysis: TenderAnalysis) -> None:
    project_name = analysis.project_info.project_name or "未命名投标项目"

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_before = Pt(76)
    title.paragraph_format.space_after = Pt(42)
    title.paragraph_format.line_spacing = 1.3
    title_size = 22 if len(project_name) <= 22 else 17 if len(project_name) <= 34 else 15.5
    title_run = title.add_run(project_name)
    _format_run(
        title_run,
        latin=COVER_FONT,
        east_asia=COVER_FONT_EAST_ASIA,
        size=title_size,
        bold=True,
        color=BLACK,
    )

    document_type = document.add_paragraph()
    document_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document_type.paragraph_format.first_line_indent = None
    document_type.paragraph_format.space_after = Pt(10)
    document_type_run = document_type.add_run("投 标 文 件")
    _format_run(
        document_type_run,
        latin=COVER_FONT,
        east_asia=COVER_FONT_EAST_ASIA,
        size=30,
        bold=True,
        color=BLACK,
        character_spacing=2,
    )

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = None
    subtitle.paragraph_format.space_after = Pt(88)
    subtitle_run = subtitle.add_run("技术及商务响应文件")
    _format_run(
        subtitle_run,
        latin=HEADING_FONT,
        east_asia=HEADING_FONT_EAST_ASIA,
        size=16,
        bold=True,
        color=BLACK,
    )

    metadata = document.add_table(rows=4, cols=2)
    _set_table_geometry(metadata, [2750, 6151], indent=0)
    _remove_table_borders(metadata)
    values = (
        ("项目名称：", project_name),
        ("招标人：", analysis.project_info.purchaser or "________________"),
        ("投标人（盖章）：", "________________"),
        ("日  期：", date.today().strftime("%Y 年 %m 月 %d 日")),
    )
    for row, row_values in zip(metadata.rows, values, strict=True):
        for index, (cell, value) in enumerate(zip(row.cells, row_values, strict=True)):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=100, bottom=100, start=80, end=80)
            if index == 0:
                _set_cell_no_wrap(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            if index == 1:
                _set_paragraph_border(paragraph, side="bottom", color=BLACK, size=6, space=1)
            _format_run(
                paragraph.runs[0],
                size=12,
                bold=index == 0,
                color=BLACK,
            )
    document.add_page_break()


def _strip_markdown(value: str) -> str:
    value = re.sub(r"\[资料\s*\d+\]", "", value)
    value = re.sub(r"\s+([，。；：！？])", r"\1", value)
    return re.sub(
        r"\*\*(.*?)\*\*|`([^`]*)`",
        lambda match: match.group(1) or match.group(2) or "",
        value,
    ).strip()


def _sanitize_markdown(markdown: str) -> str:
    lines = markdown.replace("\ufeff", "").splitlines()
    cleaned: list[str] = []
    first_content_seen = False
    in_frontmatter = False
    for line in lines:
        stripped = line.strip()
        if re.fullmatch(r"```[A-Za-z0-9_-]*", stripped):
            continue
        if not first_content_seen and not stripped:
            continue
        if not first_content_seen and stripped == "---":
            first_content_seen = True
            in_frontmatter = True
            continue
        if in_frontmatter:
            if stripped == "---":
                in_frontmatter = False
            continue
        first_content_seen = True
        if stripped == "---":
            continue
        cleaned.append(line)
    while cleaned and not cleaned[-1].strip():
        cleaned.pop()
    return "\n".join(cleaned)


def _add_list_item(document: Document, text: str, *, num_id: int, ordered: bool) -> None:
    paragraph = document.add_paragraph(style="List Number" if ordered else "List Bullet")
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.add_run(text)
    _apply_numbering(paragraph, num_id)


def _add_markdown(
    document: Document,
    markdown: str,
    *,
    bullet_num_id: int,
    chapter_title: str = "",
) -> None:
    lines = _sanitize_markdown(markdown).splitlines()
    index = 0
    last_content_kind = ""
    ordered_num_id: int | None = None
    title_skipped = False
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            heading_text = _strip_markdown(heading_match.group(2))
            if not title_skipped and chapter_title and heading_text.strip() == chapter_title.strip():
                title_skipped = True
                last_content_kind = "heading"
                index += 1
                continue
        if line.startswith("### "):
            document.add_heading(_strip_markdown(line[4:]), level=3)
            last_content_kind = "heading"
        elif line.startswith("## "):
            document.add_heading(_strip_markdown(line[3:]), level=2)
            last_content_kind = "heading"
        elif line.startswith("# "):
            document.add_heading(_strip_markdown(line[2:]), level=2)
            last_content_kind = "heading"
        elif line.startswith(("- ", "* ")):
            _add_list_item(
                document,
                _strip_markdown(line[2:]),
                num_id=bullet_num_id,
                ordered=False,
            )
            last_content_kind = "bullet"
        elif re.match(r"^\d+[.、]\s*", line):
            if last_content_kind != "ordered" or ordered_num_id is None:
                ordered_num_id = _create_numbering(document, ordered=True)
            _add_list_item(
                document,
                _strip_markdown(re.sub(r"^\d+[.、]\s*", "", line)),
                num_id=ordered_num_id,
                ordered=True,
            )
            last_content_kind = "ordered"
        elif "|" in line and index + 1 < len(lines) and re.fullmatch(
            r"\s*\|?[\s:|-]+\|?\s*", lines[index + 1]
        ):
            headers = [cell.strip() for cell in line.strip("|").split("|")]
            rows: list[list[str]] = []
            index += 2
            while index < len(lines) and "|" in lines[index]:
                rows.append([cell.strip() for cell in lines[index].strip("|").split("|")])
                index += 1
            table = document.add_table(rows=1, cols=len(headers))
            for column, value in enumerate(headers):
                table.rows[0].cells[column].text = _strip_markdown(value)
            for row in rows:
                cells = table.add_row().cells
                for column in range(len(headers)):
                    cells[column].text = _strip_markdown(row[column] if column < len(row) else "")
            _style_table(
                table,
                _column_widths(headers, rows),
                center_columns=_center_columns(headers),
                nowrap_columns=_center_columns(headers),
            )
            last_content_kind = "table"
            continue
        elif line.lstrip().startswith(">"):
            paragraph = document.add_paragraph(_strip_markdown(line.lstrip()[1:].strip()))
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.left_indent = Cm(0.65)
            paragraph.paragraph_format.right_indent = Cm(0.35)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.35
            _set_paragraph_shading(paragraph, PALE_GRAY)
            _set_paragraph_border(paragraph, side="left", color=BORDER, size=10, space=6)
            last_content_kind = "quote"
        else:
            paragraph = document.add_paragraph(_strip_markdown(line))
            paragraph.paragraph_format.first_line_indent = Cm(0.85)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(26)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            last_content_kind = "paragraph"
        index += 1


def _add_analysis_appendix(
    document: Document,
    analysis: TenderAnalysis,
    *,
    bullet_num_id: int,
) -> None:
    heading = document.add_heading("内部核对事项", level=1)
    heading.paragraph_format.page_break_before = True

    document.add_heading("强制要求", level=2)
    for item in analysis.mandatory_requirements:
        page = f"第 {item.source_page} 页" if item.source_page else "页码未知"
        _add_list_item(
            document,
            f"[{item.status}] {item.content}（{page}）",
            num_id=bullet_num_id,
            ordered=False,
        )

    document.add_heading("评分项", level=2)
    if analysis.scoring_items:
        table = document.add_table(rows=1, cols=4)
        for cell, value in zip(table.rows[0].cells, ("评分项", "分值", "页码", "状态"), strict=True):
            cell.text = value
        for item in analysis.scoring_items:
            cells = table.add_row().cells
            cells[0].text = item.criterion
            cells[1].text = item.points
            cells[2].text = str(item.source_page or "")
            cells[3].text = item.status
        _style_table(table, [4900, 1000, 1000, 2001], center_columns={1, 2, 3})
    else:
        document.add_paragraph("未识别到评分项，请人工检查招标文件。")

    document.add_heading("风险与待确认项", level=2)
    risk_items = [
        *analysis.risks,
        *[item for item in analysis.mandatory_requirements if item.status != "已确认"],
    ]
    for item in risk_items[:50]:
        _add_list_item(document, item.content, num_id=bullet_num_id, ordered=False)


def _add_review_appendix(document: Document, review: ReviewReport) -> None:
    heading = document.add_heading("内部复核报告", level=1)
    heading.paragraph_format.page_break_before = True
    summary = document.add_paragraph()
    summary.paragraph_format.first_line_indent = None
    summary.paragraph_format.space_before = Pt(2)
    summary.paragraph_format.space_after = Pt(10)
    summary.paragraph_format.line_spacing = 1.3
    _set_paragraph_shading(summary, LIGHT_GRAY)
    _set_paragraph_border(
        summary,
        side="left",
        color=BORDER,
        size=10,
        space=8,
    )
    summary_run = summary.add_run(
        f"待处理问题 {review.pending_count()} 项，其中高风险 {review.severity_count('高')} 项、"
        f"中风险 {review.severity_count('中')} 项、低风险 {review.severity_count('低')} 项。"
    )
    _format_run(
        summary_run,
        size=10,
        bold=True,
        color=BLACK,
    )
    if not review.issues:
        document.add_paragraph("当前规则未发现问题，仍需由投标负责人完成最终审核。")
        return

    table = document.add_table(rows=1, cols=4)
    for cell, value in zip(table.rows[0].cells, ("级别 / 类别", "问题与处理建议", "页码", "状态"), strict=True):
        cell.text = value
    for item in review.issues:
        cells = table.add_row().cells
        cells[0].text = f"{item.severity}\n{item.category}"
        cells[1].text = ""
        message = cells[1].paragraphs[0]
        message_run = message.add_run(item.message)
        message_run.bold = True
        suggestion = cells[1].add_paragraph()
        suggestion_run = suggestion.add_run(f"建议：{item.suggestion}")
        suggestion_run.font.color.rgb = RGBColor.from_string(MUTED_TEXT)
        cells[2].text = str(item.source_page or "")
        cells[3].text = item.status
    _style_table(
        table,
        [1500, 4500, 1200, 1701],
        center_columns={0, 2, 3},
        nowrap_columns={2, 3},
        compact=True,
    )
    for row, issue in zip(table.rows[1:], review.issues, strict=True):
        if issue.severity == "高":
            for run in row.cells[0].paragraphs[0].runs:
                _format_run(run, size=12, bold=True, color=BLACK)


def _add_contents(document: Document, drafts: list[ChapterDraft]) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = None
    heading.paragraph_format.space_before = Pt(30)
    heading.paragraph_format.space_after = Pt(26)
    run = heading.add_run("目 录")
    _format_run(
        run,
        latin=HEADING_FONT,
        east_asia=HEADING_FONT_EAST_ASIA,
        size=18,
        bold=True,
        color=BLACK,
        character_spacing=1,
    )
    for index, draft in enumerate(drafts, start=1):
        entry = document.add_paragraph()
        entry.paragraph_format.first_line_indent = None
        entry.paragraph_format.left_indent = Cm(0.8)
        entry.paragraph_format.right_indent = Cm(0.8)
        entry.paragraph_format.space_after = Pt(8)
        entry.paragraph_format.line_spacing = 1.4
        entry_run = entry.add_run(f"第 {index} 章  {draft.title}")
        _format_run(entry_run, size=12, color=BLACK)
    document.add_page_break()


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            yield from container.paragraphs
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def _replace_paragraph_text(paragraph, placeholder: str, value: str) -> None:
    if placeholder not in paragraph.text:
        return
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            return
    start = paragraph.text.find(placeholder)
    end = start + len(placeholder)
    offset = 0
    started = False
    for run in paragraph.runs:
        run_start = offset
        run_end = offset + len(run.text)
        offset = run_end
        if run_end <= start or run_start >= end:
            continue
        prefix = run.text[: max(0, start - run_start)] if not started else ""
        suffix = run.text[max(0, end - run_start) :] if run_end >= end else ""
        run.text = f"{prefix}{value if not started else ''}{suffix}"
        started = True


def _apply_template_placeholders(document: Document, analysis: TenderAnalysis) -> None:
    info = analysis.project_info
    replacements = {
        "{{PROJECT_NAME}}": info.project_name,
        "{{PURCHASER}}": info.purchaser,
        "{{AGENCY}}": info.agency,
        "{{BUDGET}}": info.budget,
        "{{BID_DEADLINE}}": info.bid_deadline,
        "{{GENERATED_DATE}}": date.today().isoformat(),
    }
    for paragraph in _iter_all_paragraphs(document):
        for placeholder, value in replacements.items():
            _replace_paragraph_text(paragraph, placeholder, value or "")


def _add_qualification_images(document: Document, image_paths: list[Path]) -> None:
    valid_paths = [Path(path) for path in image_paths if Path(path).is_file()]
    if not valid_paths:
        return
    heading = document.add_heading("资质证明材料", level=1)
    heading.paragraph_format.page_break_before = True
    intro = document.add_paragraph("本章按材料顺序编排资质证明文件。")
    intro.paragraph_format.first_line_indent = None
    intro.paragraph_format.space_after = Pt(12)

    section = document.sections[-1]
    usable_width_dxa = max(
        3000,
        int((section.page_width - section.left_margin - section.right_margin) / 635),
    )
    first_column = usable_width_dxa // 2
    table = document.add_table(rows=(len(valid_paths) + 1) // 2, cols=2)
    _style_table(table, [first_column, usable_width_dxa - first_column], compact=True)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_shading(cell, "FFFFFF")
    for index, path in enumerate(valid_paths):
        cell = table.cell(index // 2, index % 2)
        cell.text = ""
        picture_paragraph = cell.paragraphs[0]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.first_line_indent = None
        picture_paragraph.paragraph_format.space_after = Pt(5)
        with Image.open(path) as image:
            width_px, height_px = image.size
        max_width_cm = min(7.0, max(3.0, first_column * 2.54 / 1440 - 0.8))
        scale = min(max_width_cm / max(width_px, 1), 9.8 / max(height_px, 1))
        width_cm = max(1.0, width_px * scale)
        height_cm = max(1.0, height_px * scale)
        picture_paragraph.add_run().add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
        caption = cell.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = None
        caption.paragraph_format.space_after = Pt(4)
        caption_text = re.sub(r"[_-]+", " ", path.stem).strip() or f"资质材料 {index + 1}"
        _format_run(caption.add_run(caption_text), size=12, color=BLACK)


def _add_generated_sections(
    document: Document,
    analysis: TenderAnalysis,
    drafts: list[ChapterDraft],
    review: ReviewReport | None,
    qualification_images: list[Path],
    *,
    include_contents: bool,
    include_internal_appendices: bool,
) -> None:
    bullet_num_id = _create_numbering(document, ordered=False)
    if include_contents:
        _add_contents(document, drafts)
    for index, draft in enumerate(drafts, start=1):
        heading = document.add_heading(f"第 {index} 章 {draft.title}", level=1)
        if index > 1:
            heading.paragraph_format.page_break_before = True
        _add_markdown(
            document,
            draft.markdown,
            bullet_num_id=bullet_num_id,
            chapter_title=draft.title,
        )
    if include_internal_appendices:
        _add_analysis_appendix(document, analysis, bullet_num_id=bullet_num_id)
        if review is not None:
            _add_review_appendix(document, review)
    _add_qualification_images(document, qualification_images)


def _insert_template_content(
    document: Document,
    analysis: TenderAnalysis,
    drafts: list[ChapterDraft],
    review: ReviewReport | None,
    qualification_images: list[Path],
    *,
    include_internal_appendices: bool,
) -> None:
    placeholder = next(
        (paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "{{BID_CONTENT}}"),
        None,
    )
    body = document._element.body
    existing = set(body.iterchildren())
    _add_generated_sections(
        document,
        analysis,
        drafts,
        review,
        qualification_images,
        include_contents=True,
        include_internal_appendices=include_internal_appendices,
    )
    if placeholder is None:
        return
    new_elements = [
        element for element in body.iterchildren()
        if element not in existing and element.tag != qn("w:sectPr")
    ]
    for element in new_elements:
        placeholder._p.addprevious(element)
    body.remove(placeholder._p)


def export_docx(
    output_path: str | Path,
    analysis: TenderAnalysis,
    drafts: list[ChapterDraft],
    review: ReviewReport | None = None,
    *,
    template_path: str | Path | None = None,
    qualification_images: list[str | Path] | None = None,
    include_internal_appendices: bool = False,
) -> Path:
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    images = [Path(path) for path in qualification_images or []]
    document = Document(str(template_path)) if template_path else Document()
    if not template_path:
        _configure_document(document)
    document.core_properties.title = analysis.project_info.project_name or "投标文件初稿"
    document.core_properties.subject = "投标响应文件"
    document.core_properties.author = ""

    if template_path:
        _apply_template_placeholders(document, analysis)
        _insert_template_content(
            document,
            analysis,
            drafts,
            review,
            images,
            include_internal_appendices=include_internal_appendices,
        )
    else:
        _add_cover(document, analysis)
        _add_contents(document, drafts)
        _add_generated_sections(
            document,
            analysis,
            drafts,
            review,
            images,
            include_contents=False,
            include_internal_appendices=include_internal_appendices,
        )
        _add_header_footer(document, analysis.project_info.project_name or "未命名投标项目")
    document.save(target)
    return target
